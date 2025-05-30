from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
import os

# Ensure test_files directory exists
os.makedirs('test_files', exist_ok=True)

def create_test_pdf():
    """Create a sample PDF file with some text and metadata for testing"""
    print("Creating sample PDF file...")
    
    # Create PDF with reportlab
    pdf_path = 'test_files/sample_test.pdf'
    doc = SimpleDocTemplate(pdf_path, pagesize=letter)
    styles = getSampleStyleSheet()
    
    # Create content
    content = []
    
    # Add title
    content.append(Paragraph("Sample Test Document", styles['Title']))
    content.append(Spacer(1, 0.25*inch))
    
    # Add some paragraphs
    for i in range(5):
        text = f"This is paragraph {i+1} of the sample document. " * 5
        content.append(Paragraph(text, styles['Normal']))
        content.append(Spacer(1, 0.2*inch))
    
    # Add a hyperlink
    link_text = "This is a sample <a href='https://www.python.org' color='blue'>hyperlink</a> to the Python website."
    content.append(Paragraph(link_text, styles['Normal']))
    content.append(Spacer(1, 0.2*inch))
    
    # Add more paragraphs
    for i in range(5, 10):
        text = f"This is paragraph {i+1} of the sample document with some different content. " * 3
        content.append(Paragraph(text, styles['Normal']))
        content.append(Spacer(1, 0.2*inch))
    
    # Build the PDF
    doc.build(content)
    
    # Add metadata using pyPDF
    try:
        from pypdf import PdfWriter, PdfReader
        
        # Read the PDF we just created
        reader = PdfReader(pdf_path)
        writer = PdfWriter()
        
        # Add all pages from reader to writer
        for page in reader.pages:
            writer.add_page(page)
        
        # Add metadata
        writer.add_metadata({
            "/Author": "SafeFileGuard Testing",
            "/Subject": "Sample Document for Testing",
            "/Title": "PDF Analysis Test",
            "/Keywords": "pdf,test,analysis,sample,document"
        })
        
        # Save the PDF with metadata
        with open(pdf_path, "wb") as output_file:
            writer.write(output_file)
        
        print(f"Successfully created test PDF with metadata at {pdf_path}")
    except Exception as e:
        print(f"Error adding metadata: {str(e)}")
        print(f"Basic PDF was still created at {pdf_path}")

def create_test_docx():
    """Create a sample DOCX file with some text and metadata for testing"""
    try:
        import docx
        from docx.shared import Inches
        from docx.enum.dml import MSO_THEME_COLOR_INDEX
        
        print("Creating sample DOCX file...")
        
        doc = docx.Document()
        
        # Add document properties
        doc.core_properties.author = "SafeFileGuard Testing"
        doc.core_properties.title = "Word Document Analysis Test"
        doc.core_properties.subject = "Sample Document for Testing"
        doc.core_properties.keywords = "docx,test,analysis,sample,document"
        
        # Add heading
        doc.add_heading('Sample Test Document', 0)
        
        # Add some paragraphs
        for i in range(10):
            doc.add_paragraph(f"This is paragraph {i+1} of the sample document. " * 3)
        
        # Add a table
        table = doc.add_table(rows=3, cols=3)
        for row in range(3):
            for col in range(3):
                cell = table.cell(row, col)
                cell.text = f"Row {row+1}, Column {col+1}"
        
        # Add a hyperlink
        paragraph = doc.add_paragraph("Here is a link to ")
        run = paragraph.add_run("Python's website")
        
        # Adding hyperlink
        r_id = doc.part.relate_to('https://www.python.org', docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)
        hyperlink.append(run._element)
        paragraph._p.append(hyperlink)
        
        # Save the document
        docx_path = 'test_files/sample_test.docx'
        doc.save(docx_path)
        
        print(f"Successfully created test DOCX at {docx_path}")
    except Exception as e:
        print(f"Error creating DOCX: {str(e)}")

if __name__ == "__main__":
    create_test_pdf()
    create_test_docx()
    print("Test files created successfully.") 