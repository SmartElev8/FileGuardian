from flask import Flask, request, jsonify, render_template, redirect, url_for, session
import os
from werkzeug.utils import secure_filename
import subprocess
import json
import datetime
import mimetypes
import re
from pathlib import Path
import io
import logging
import vt
from dotenv import load_dotenv
import uuid
import shutil
import sqlite3
from flask_sqlalchemy import SQLAlchemy
import platform
import hashlib
import time

# Load environment variables
load_dotenv()

# PDF libraries
try:
    from pypdf import PdfReader
    import pikepdf
except ImportError:
    logging.warning("PDF libraries not available. PDF analysis will be limited.")

# DOCX libraries
try:
    import docx
except ImportError:
    logging.warning("python-docx not available. DOCX analysis will be limited.")

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.config['ALLOWED_EXTENSIONS'] = {'exe', 'pdf', 'docx', 'doc', 'ppt', 'pptx', 'xls', 'xlsx', 'txt', 'rtf', 'odt', 'zip', 'rar', 'tar', 'gz'}
app.config['QUARANTINE_FOLDER'] = 'quarantine'
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'smartfileguardian-secret-key')
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///smartfileguardian.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# Initialize SQLAlchemy
db = SQLAlchemy(app)

# Database Models
class QuarantinedFile(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    device_id = db.Column(db.String(64), nullable=False)
    file_name = db.Column(db.String(255), nullable=False)
    original_path = db.Column(db.String(512))
    quarantine_path = db.Column(db.String(512), nullable=False)
    file_size = db.Column(db.Integer)
    file_type = db.Column(db.String(64))
    date_quarantined = db.Column(db.DateTime, default=datetime.datetime.utcnow)
    is_malicious = db.Column(db.Boolean, default=True)
    threat_details = db.Column(db.Text)
    
    def to_dict(self):
        return {
            'id': self.id,
            'file_name': self.file_name,
            'original_path': self.original_path,
            'file_size': self.file_size,
            'file_type': self.file_type,
            'date_quarantined': self.date_quarantined.strftime('%Y-%m-%d %H:%M:%S'),
            'is_malicious': self.is_malicious,
            'threat_details': self.threat_details
        }

class ScannedFile(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    device_id = db.Column(db.String(64), nullable=False)
    file_name = db.Column(db.String(255), nullable=False)
    file_path = db.Column(db.String(512))
    file_size = db.Column(db.Integer)
    file_type = db.Column(db.String(64))
    date_scanned = db.Column(db.DateTime, default=datetime.datetime.utcnow)
    is_malicious = db.Column(db.Boolean, default=False)
    scan_result = db.Column(db.Text)
    
    def to_dict(self):
        return {
            'id': self.id,
            'file_name': self.file_name,
            'file_path': self.file_path,
            'file_size': self.file_size,
            'file_type': self.file_type,
            'date_scanned': self.date_scanned.strftime('%Y-%m-%d %H:%M:%S'),
            'is_malicious': self.is_malicious,
            'scan_result': self.scan_result
        }

class UserActivity(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    device_id = db.Column(db.String(64), nullable=False)
    activity_type = db.Column(db.String(64), nullable=False)
    description = db.Column(db.Text)
    file_name = db.Column(db.String(255))
    timestamp = db.Column(db.DateTime, default=datetime.datetime.utcnow)
    
    def to_dict(self):
        return {
            'id': self.id,
            'activity_type': self.activity_type,
            'description': self.description,
            'file_name': self.file_name,
            'timestamp': self.timestamp.strftime('%Y-%m-%d %H:%M:%S')
        }

# Create database tables
with app.app_context():
    db.create_all()

# Ensure upload and quarantine directories exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['QUARANTINE_FOLDER'], exist_ok=True)

# Ensure static directory exists
os.makedirs(os.path.join('static', 'images'), exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

def get_device_id():
    """Generate or retrieve a unique device identifier"""
    if 'device_id' not in session:
        # Generate a unique device ID based on hardware info
        system_info = platform.system() + platform.node() + platform.machine()
        # Add timestamp to make it more unique
        system_info += str(time.time())
        # Create a hash of the system info
        device_hash = hashlib.md5(system_info.encode()).hexdigest()
        session['device_id'] = device_hash
    return session['device_id']

def quarantine_file(file_path, original_path=None, is_malicious=True, threat_details=None):
    """Move a file to quarantine and record in database"""
    try:
        # Generate unique filename for quarantine
        filename = os.path.basename(file_path)
        file_extension = os.path.splitext(filename)[1]
        quarantine_filename = f"{uuid.uuid4().hex}{file_extension}"
        quarantine_path = os.path.join(app.config['QUARANTINE_FOLDER'], quarantine_filename)
        
        # Copy file to quarantine
        shutil.copy2(file_path, quarantine_path)
        
        # Get file details
        file_size = os.path.getsize(file_path)
        file_type = mimetypes.guess_type(file_path)[0] or 'application/octet-stream'
        
        # Store in database
        device_id = get_device_id()
        quarantined_file = QuarantinedFile(
            device_id=device_id,
            file_name=filename,
            original_path=original_path or file_path,
            quarantine_path=quarantine_path,
            file_size=file_size,
            file_type=file_type,
            is_malicious=is_malicious,
            threat_details=json.dumps(threat_details) if threat_details else None
        )
        db.session.add(quarantined_file)
        
        # Log activity
        activity = UserActivity(
            device_id=device_id,
            activity_type='quarantine',
            description=f"File '{filename}' was quarantined",
            file_name=filename
        )
        db.session.add(activity)
        db.session.commit()
        
        return True, quarantined_file.id
    except Exception as e:
        logger.error(f"Error quarantining file: {str(e)}")
        return False, str(e)

def restore_file(quarantine_id, restore_path=None):
    """Restore a file from quarantine"""
    try:
        # Get quarantined file record
        quarantined_file = QuarantinedFile.query.get(quarantine_id)
        if not quarantined_file:
            return False, "File not found in quarantine"
        
        # Determine restore path
        if restore_path:
            # Use provided path
            target_path = os.path.join(restore_path, quarantined_file.file_name)
        elif quarantined_file.original_path and os.path.dirname(quarantined_file.original_path):
            # Use original path if it exists
            target_path = quarantined_file.original_path
        else:
            # Default to desktop
            target_path = os.path.join(os.path.expanduser("~"), "Desktop", quarantined_file.file_name)
        
        # Ensure target directory exists
        os.makedirs(os.path.dirname(target_path), exist_ok=True)
        
        # Copy file from quarantine to target path
        shutil.copy2(quarantined_file.quarantine_path, target_path)
        
        # Log activity
        device_id = get_device_id()
        activity = UserActivity(
            device_id=device_id,
            activity_type='restore',
            description=f"File '{quarantined_file.file_name}' was restored to {target_path}",
            file_name=quarantined_file.file_name
        )
        db.session.add(activity)
        db.session.commit()
        
        return True, target_path
    except Exception as e:
        logger.error(f"Error restoring file: {str(e)}")
        return False, str(e)

def delete_quarantined_file(quarantine_id):
    """Delete a file from quarantine"""
    try:
        # Get quarantined file record
        quarantined_file = QuarantinedFile.query.get(quarantine_id)
        if not quarantined_file:
            return False, "File not found in quarantine"
        
        # Delete the physical file
        if os.path.exists(quarantined_file.quarantine_path):
            os.remove(quarantined_file.quarantine_path)
        
        # Log activity
        device_id = get_device_id()
        activity = UserActivity(
            device_id=device_id,
            activity_type='delete',
            description=f"File '{quarantined_file.file_name}' was deleted from quarantine",
            file_name=quarantined_file.file_name
        )
        db.session.add(activity)
        
        # Delete the database record
        db.session.delete(quarantined_file)
        db.session.commit()
        
        return True, "File deleted successfully"
    except Exception as e:
        logger.error(f"Error deleting quarantined file: {str(e)}")
        return False, str(e)

def scan_with_virustotal(file_path=None, url=None):
    """Scan a file or URL using VirusTotal API"""
    try:
        api_key = os.getenv('VT_API_KEY')
        if not api_key:
            raise ValueError("VirusTotal API key not found in environment variables")

        client = vt.Client(api_key)
        results = {}

        if file_path:
            with open(file_path, 'rb') as f:
                analysis = client.scan_file(f)
            analysis_id = analysis.id
            analysis_result = client.get_object(f"/analyses/{analysis_id}")
            results = {
                'id': analysis_id,
                'status': analysis_result.status,
                'stats': getattr(analysis_result, 'stats', None)
            }
        elif url:
            analysis = client.scan_url(url)
            analysis_id = analysis.id
            analysis_result = client.get_object(f"/analyses/{analysis_id}")
            results = {
                'id': analysis_id,
                'status': analysis_result.status,
                'url': url
            }

        client.close()
        return results
    except Exception as e:
        logger.error(f"Error in VirusTotal scan: {str(e)}")
        return None

def validate_ml_models():
    """Validate that all required ML models are present"""
    required_models = [
        'Classifier/classifier.pkl',
        'Classifier/features.pkl', 
        'Classifier/pickel_model.pkl',
        'Classifier/pickel_vector.pkl'
    ]
    
    missing_models = []
    for model_path in required_models:
        if not os.path.exists(model_path):
            missing_models.append(model_path)
    
    if missing_models:
        logger.error(f"Missing required ML models: {missing_models}")
        return False, missing_models
    
    return True, []

def run_local_file_scan(filepath: str) -> dict:
    """Run local file scanning using existing models"""
    try:
        # Validate models first
        models_valid, missing_models = validate_ml_models()
        if not models_valid:
            return {
                'is_malicious': False,
                'message': f'Error: Missing ML models: {missing_models}. Please ensure all trained models are present in Classifier/ directory.',
                'details': {'error': 'Missing ML models', 'missing': missing_models},
                'scan_type': 'local'
            }
        
        # Determine which scanner to use based on file extension
        file_extension = os.path.splitext(filepath)[1].lower()
        if file_extension in ['.pdf', '.docx']:
            scanner_script = 'Extract/document_scanner/document_main.py'
        else:
            scanner_script = 'Extract/PE_main.py'
        
        # Run the appropriate scanner
        result = subprocess.run(['python', scanner_script, filepath], 
                              capture_output=True, text=True)
        
        # Parse the result
        try:
            result_data = json.loads(result.stdout)
            is_malicious = result_data.get('is_malicious', False)
            message = result_data.get('message', 'Analysis completed')
            details = result_data.get('details', {})
        except json.JSONDecodeError:
            is_malicious = 'malicious' in result.stdout.lower()
            message = result.stdout
            details = {}
        
        return {
            'is_malicious': is_malicious,
            'message': message,
            'details': details,
            'scan_type': 'local'
        }
    except Exception as e:
        return {
            'is_malicious': False,
            'message': f'Error during local analysis: {str(e)}',
            'details': {'error': str(e)},
            'scan_type': 'local'
        }

def run_local_url_scan(url: str) -> dict:
    """Run local URL scanning using existing models"""
    try:
        result = subprocess.run(['python', 'Extract/url_main.py', url], 
                              capture_output=True, text=True)
        
        try:
            result_data = json.loads(result.stdout)
            is_malicious = result_data.get('is_malicious', False)
            message = result_data.get('message', 'Analysis completed')
            details = result_data.get('details', {})
        except json.JSONDecodeError:
            is_malicious = 'malicious' in result.stdout.lower()
            message = result.stdout
            details = {}
        
        return {
            'is_malicious': is_malicious,
            'message': message,
            'details': details,
            'scan_type': 'local'
        }
    except Exception as e:
        return {
            'is_malicious': False,
            'message': f'Error during local analysis: {str(e)}',
            'details': {'error': str(e)},
            'scan_type': 'local'
        }

@app.route('/')
def home():
    return render_template('index.html')

@app.route('/documentation')
def documentation():
    return render_template('documentation.html')

@app.route('/api')
def api_docs():
    return render_template('api.html')

@app.route('/api-reference')
def api_reference():
    return render_template('api_reference.html')

@app.route('/faqs')
def faqs():
    return render_template('faqs.html')

@app.route('/threat-database')
def threat_database():
    return render_template('threat-database.html')

@app.route('/support')
def support():
    return render_template('support.html')

def extract_text_from_pdf(pdf_path, max_pages=10):
    """Extract text from the first few pages of a PDF for summarization"""
    try:
        text = ""
        with open(pdf_path, 'rb') as file:
            reader = PdfReader(file)
            num_pages = min(len(reader.pages), max_pages)
            
            for i in range(num_pages):
                page = reader.pages[i]
                text += page.extract_text() + "\n"
                
        return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        return ""

def count_links_in_pdf(pdf_path):
    """Count links in a PDF file"""
    try:
        with pikepdf.open(pdf_path) as pdf:
            link_count = 0
            for page in pdf.pages:
                if "/Annots" in page:
                    annotations = page["/Annots"]
                    for annot in annotations:
                        annot_obj = annot.get_object()
                        if annot_obj.get("/Subtype") == "/Link":
                            link_count += 1
            return link_count
    except Exception as e:
        logger.error(f"Error counting links in PDF: {str(e)}")
        return 0

def extract_text_from_docx(docx_path):
    """Extract text from a DOCX file for summarization"""
    try:
        doc = docx.Document(docx_path)
        text = ""
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            text += para.text + "\n"
            
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
                
        return text
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        return ""

def count_links_in_docx(docx_path):
    """Count hyperlinks in a DOCX file"""
    try:
        doc = docx.Document(docx_path)
        link_count = 0
        
        # Count hyperlinks in paragraphs
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.element.findall('.//w:hyperlink', namespaces=docx.oxml.ns.nsmap):
                    link_count += 1
        
        return link_count
    except Exception as e:
        logger.error(f"Error counting links in DOCX: {str(e)}")
        return 0

def count_images(file_path, file_type):
    """Count images in document files"""
    try:
        if file_type == '.pdf':
            # Count images in PDF
            with pikepdf.open(file_path) as pdf:
                image_count = 0
                for page in pdf.pages:
                    if '/Resources' in page and '/XObject' in page['/Resources']:
                        xobjects = page['/Resources']['/XObject']
                        for obj in xobjects:
                            if xobjects[obj].get('/Subtype') == '/Image':
                                image_count += 1
                return image_count
                
        elif file_type in ['.doc', '.docx']:
            # Count images in DOCX
            doc = docx.Document(file_path)
            image_count = 0
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_count += 1
            return image_count
            
        return 0
    except Exception as e:
        logger.error(f"Error counting images: {str(e)}")
        return 0

def generate_content_summary(text, max_length=200):
    """Generate a brief summary of the document content"""
    try:
        # Simple summary: first few sentences up to max_length
        text = text.strip()
        sentences = re.split(r'(?<=[.!?])\s+', text)
        summary = ""
        
        for sentence in sentences:
            if len(summary) + len(sentence) <= max_length:
                summary += sentence + " "
            else:
                break
                
        if summary:
            return summary.strip()
        else:
            return "This document appears to contain no extractable text content."
    except Exception as e:
        logger.error(f"Error generating summary: {str(e)}")
        return "Unable to generate content summary."

def get_file_details(filepath):
    """Extract detailed information from files using appropriate libraries"""
    try:
        file_path = Path(filepath)
        file_stats = os.stat(filepath)
        file_size = file_stats.st_size
        
        # Use mimetypes to get file type
        mime_type = mimetypes.guess_type(filepath)[0] or 'application/octet-stream'
        
        # Get basic file information
        details = {
            'file_size': file_size,
            'file_type': mime_type,
            'created_date': datetime.datetime.fromtimestamp(file_stats.st_ctime).strftime('%Y-%m-%d %H:%M:%S'),
            'modified_date': datetime.datetime.fromtimestamp(file_stats.st_mtime).strftime('%Y-%m-%d %H:%M:%S'),
        }
        
        # Handle different file types
        file_extension = file_path.suffix.lower()
        
        # PDF Files
        if file_extension == '.pdf' or mime_type == 'application/pdf':
            try:
                with open(filepath, 'rb') as f:
                    pdf = PdfReader(f)
                    
                    # Extract metadata
                    metadata = pdf.metadata
                    if metadata:
                        if metadata.get('/Author'):
                            details['author'] = metadata.get('/Author')
                        if metadata.get('/Keywords'):
                            details['keywords'] = metadata.get('/Keywords')
                        if metadata.get('/Subject'):
                            details['subject'] = metadata.get('/Subject')
                        if metadata.get('/Title'):
                            details['title'] = metadata.get('/Title')
                    
                    # Page count
                    details['page_count'] = len(pdf.pages)
                    
                    # Count links
                    details['links_count'] = count_links_in_pdf(filepath)
                    
                    # Count images/possible QR codes
                    details['image_count'] = count_images(filepath, '.pdf')
                    details['qr_codes_count'] = 'Unknown' # Would need specific QR detection library
                    
                    # Extract text for summary
                    text = extract_text_from_pdf(filepath)
                    details['content_summary'] = generate_content_summary(text)
                    
            except Exception as e:
                logger.error(f"Error analyzing PDF: {str(e)}")
                details.update({
                    'page_count': 'Error',
                    'content_summary': f"Could not analyze PDF content: {str(e)}"
                })
        
        # DOC/DOCX Files
        elif file_extension in ['.doc', '.docx'] or mime_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
            try:
                doc = docx.Document(filepath)
                
                # Basic document properties
                core_properties = doc.core_properties
                if core_properties.author:
                    details['author'] = core_properties.author
                if core_properties.keywords:
                    details['keywords'] = core_properties.keywords
                if core_properties.subject:
                    details['subject'] = core_properties.subject
                if core_properties.title:
                    details['title'] = core_properties.title
                
                # Page count (estimate based on paragraphs)
                # DOCX doesn't store page count directly, this is an estimation
                para_count = len(doc.paragraphs)
                details['page_count'] = max(1, para_count // 40)  # Rough estimate: ~40 paragraphs per page
                
                # Count links
                details['links_count'] = count_links_in_docx(filepath)
                
                # Count images
                details['image_count'] = count_images(filepath, '.docx')
                details['qr_codes_count'] = 'Unknown'
                
                # Extract text for summary
                text = extract_text_from_docx(filepath)
                details['content_summary'] = generate_content_summary(text)
                
            except Exception as e:
                logger.error(f"Error analyzing DOCX: {str(e)}")
                details.update({
                    'page_count': 'Error',
                    'content_summary': f"Could not analyze document content: {str(e)}"
                })
        
        # For all other file types, provide basic info
        else:
            # Extract info based on filename
            filename = file_path.stem
            words = re.findall(r'\w+', filename)
            title_case_words = [word.capitalize() for word in words]
            readable_name = ' '.join(title_case_words)
            
            details.update({
                'page_count': 'N/A',
                'links_count': 'N/A',
                'image_count': 'N/A',
                'qr_codes_count': 'N/A',
                'content_summary': f'This file appears to be a {file_extension[1:].upper()} file named "{readable_name}". The file was scanned for malicious content and found to be safe.'
            })
        
        return details
        
    except Exception as e:
        logger.error(f"Error in file analysis: {str(e)}")
        # Return basic information in case of error
        return {
            'file_size': os.path.getsize(filepath) if os.path.exists(filepath) else 0,
            'file_type': 'Unknown',
            'error': str(e),
            'content_summary': 'Error analyzing file content'
        }

@app.route('/scan/file', methods=['POST'])
def scan_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)
    
    try:
        # Get device ID
        device_id = get_device_id()
        original_path = request.form.get('original_path', '')
        
        # First try VirusTotal scan
        vt_result = scan_with_virustotal(file_path=filepath)
        
        if vt_result:
            # Process VirusTotal results
            is_malicious = vt_result.get('stats', {}).get('malicious', 0) > 0
            message = 'File is potentially malicious' if is_malicious else 'File appears to be safe'
            details = {
                'scan_type': 'virustotal',
                'stats': vt_result.get('stats', {}),
                'status': vt_result.get('status')
            }
        else:
            # Fallback to local scan
            logger.info("Falling back to local scan")
            result = run_local_file_scan(filepath)
            is_malicious = result.get('is_malicious', False)
            message = result.get('message', 'File analysis completed')
            details = result.get('details', {})
            details['scan_type'] = 'local'
        
        # Get additional file details
        file_details = get_file_details(filepath)
        logger.info(f"File details: {file_details}")
        
        # Add file details to the result
        if 'details' not in details:
            details = {}
        details.update(file_details)
        
        # Record the scan in the database
        scanned_file = ScannedFile(
            device_id=device_id,
            file_name=filename,
            file_path=original_path or filepath,
            file_size=file_details.get('file_size', 0),
            file_type=file_details.get('file_type', 'Unknown'),
            is_malicious=is_malicious,
            scan_result=json.dumps(details)
        )
        db.session.add(scanned_file)
        
        # Log activity
        activity = UserActivity(
            device_id=device_id,
            activity_type='scan',
            description=f"File '{filename}' was scanned and found to be {'malicious' if is_malicious else 'safe'}",
            file_name=filename
        )
        db.session.add(activity)
        db.session.commit()
        
        # Prepare the response
        response = {
            'is_malicious': is_malicious,
            'message': message,
            'details': details,
            'file_name': filename,
            'original_path': original_path,
            'scan_id': scanned_file.id,
            'can_quarantine': True if is_malicious else False
        }
        
        # Don't delete the file if it's malicious - we'll keep it for quarantine option
        if not is_malicious:
            os.remove(filepath)
        
        return jsonify(response)
    except Exception as e:
        logger.error(f"Error in scan_file: {str(e)}")
        if os.path.exists(filepath):
            os.remove(filepath)
            
        return jsonify({
            'is_malicious': False,
            'message': f'Error during analysis: {str(e)}',
            'details': {'error': str(e)},
            'file_name': filename
        }), 500

@app.route('/scan/url', methods=['POST'])
def scan_url():
    url = request.form.get('url')
    if not url:
        return jsonify({'error': 'No URL provided'}), 400
    
    try:
        # First try VirusTotal scan
        vt_result = scan_with_virustotal(url=url)
        
        if vt_result:
            # Process VirusTotal results
            is_malicious = vt_result.get('stats', {}).get('malicious', 0) > 0
            message = 'URL is potentially malicious' if is_malicious else 'URL appears to be safe'
            details = {
                'scan_type': 'virustotal',
                'stats': vt_result.get('stats', {}),
                'status': vt_result.get('status')
            }
        else:
            # Fallback to local scan
            logger.info("Falling back to local scan")
            result = run_local_url_scan(url)
            is_malicious = result.get('is_malicious', False)
            message = result.get('message', 'URL analysis completed')
            details = result.get('details', {})
            details['scan_type'] = 'local'
        
        return jsonify({
            'is_malicious': is_malicious,
            'message': message,
            'details': details,
            'url': url
        })
    except Exception as e:
        return jsonify({
            'is_malicious': False,
            'message': f'Error during analysis: {str(e)}',
            'details': {'error': str(e)},
            'url': url
        }), 500

@app.route('/quarantine', methods=['GET'])
def quarantine_list():
    """Display list of quarantined files"""
    device_id = get_device_id()
    quarantined_files = QuarantinedFile.query.filter_by(device_id=device_id).order_by(QuarantinedFile.date_quarantined.desc()).all()
    return render_template('quarantine.html', quarantined_files=quarantined_files)

@app.route('/quarantine/file/<int:scan_id>', methods=['POST'])
def quarantine_scan_file(scan_id):
    """Quarantine a file from scan results"""
    try:
        # Get the scan record
        scanned_file = ScannedFile.query.get_or_404(scan_id)
        
        # Check if the file is from this device
        if scanned_file.device_id != get_device_id():
            return jsonify({'success': False, 'message': 'Access denied'}), 403
        
        # Get the file path from the upload directory
        filename = secure_filename(scanned_file.file_name)
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        
        if not os.path.exists(filepath):
            return jsonify({'success': False, 'message': 'File not found'}), 404
        
        # Get threat details from scan result
        threat_details = json.loads(scanned_file.scan_result) if scanned_file.scan_result else None
        
        # Move to quarantine
        success, result = quarantine_file(
            filepath, 
            original_path=scanned_file.file_path,
            is_malicious=scanned_file.is_malicious,
            threat_details=threat_details
        )
        
        if success:
            # Delete the file from uploads
            if os.path.exists(filepath):
                os.remove(filepath)
            
            return jsonify({
                'success': True, 
                'message': 'File moved to quarantine successfully',
                'quarantine_id': result
            })
        else:
            return jsonify({'success': False, 'message': f'Error: {result}'}), 500
            
    except Exception as e:
        logger.error(f"Error quarantining file: {str(e)}")
        return jsonify({'success': False, 'message': f'Error: {str(e)}'}), 500

@app.route('/quarantine/restore/<int:quarantine_id>', methods=['POST'])
def restore_quarantined_file(quarantine_id):
    """Restore a file from quarantine"""
    try:
        # Get the quarantine record
        quarantined_file = QuarantinedFile.query.get_or_404(quarantine_id)
        
        # Check if the file is from this device
        if quarantined_file.device_id != get_device_id():
            return jsonify({'success': False, 'message': 'Access denied'}), 403
        
        # Get restore path from request if provided
        restore_path = request.form.get('restore_path', None)
        
        # Restore the file
        success, result = restore_file(quarantine_id, restore_path)
        
        if success:
            return jsonify({
                'success': True, 
                'message': 'File restored successfully',
                'restore_path': result
            })
        else:
            return jsonify({'success': False, 'message': f'Error: {result}'}), 500
            
    except Exception as e:
        logger.error(f"Error restoring file: {str(e)}")
        return jsonify({'success': False, 'message': f'Error: {str(e)}'}), 500

@app.route('/quarantine/delete/<int:quarantine_id>', methods=['POST'])
def delete_from_quarantine(quarantine_id):
    """Delete a file from quarantine"""
    try:
        # Get the quarantine record
        quarantined_file = QuarantinedFile.query.get_or_404(quarantine_id)
        
        # Check if the file is from this device
        if quarantined_file.device_id != get_device_id():
            return jsonify({'success': False, 'message': 'Access denied'}), 403
        
        # Delete the file
        success, result = delete_quarantined_file(quarantine_id)
        
        if success:
            return jsonify({
                'success': True, 
                'message': 'File deleted successfully'
            })
        else:
            return jsonify({'success': False, 'message': f'Error: {result}'}), 500
            
    except Exception as e:
        logger.error(f"Error deleting file: {str(e)}")
        return jsonify({'success': False, 'message': f'Error: {str(e)}'}), 500

@app.route('/activity', methods=['GET'])
def activity_history():
    """Display user activity history"""
    device_id = get_device_id()
    activities = UserActivity.query.filter_by(device_id=device_id).order_by(UserActivity.timestamp.desc()).all()
    return render_template('activity.html', activities=activities)

@app.route('/scans', methods=['GET'])
def scan_history():
    """Display scan history"""
    device_id = get_device_id()
    scans = ScannedFile.query.filter_by(device_id=device_id).order_by(ScannedFile.date_scanned.desc()).all()
    return render_template('scans.html', scans=scans)

if __name__ == '__main__':
    app.run(host='127.0.0.1', port=5001, debug=True) 